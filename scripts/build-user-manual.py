from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Chayapon-Works-User-Manual-TH.docx"
BLUE = "0B4DDB"
DARK = "101828"
MUTED = "667085"
LIGHT_BLUE = "EAF1FF"
PALE = "F8FAFC"
GREEN = "067647"
AMBER = "B54708"
RED = "B42318"
WHITE = "FFFFFF"
FONT = "Leelawadee UI"
TABLE_WIDTH = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font_run(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, size, color=DARK, bold=False, before=0, after=6, line=1.25):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def add_field(paragraph, instruction):
    def field_run(child):
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field_run(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    field_run(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    field_run(separate)
    text = OxmlElement("w:t")
    text.text = "1"
    field_run(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run(end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style(doc.styles["Normal"], 10.5, DARK, False, 0, 6, 1.25)
    set_style(doc.styles["Title"], 30, DARK, True, 0, 8, 1.0)
    set_style(doc.styles["Subtitle"], 13, MUTED, False, 0, 10, 1.15)
    set_style(doc.styles["Heading 1"], 16, BLUE, True, 18, 10, 1.15)
    set_style(doc.styles["Heading 2"], 13, BLUE, True, 14, 7, 1.15)
    set_style(doc.styles["Heading 3"], 11.5, "1F4D78", True, 10, 5, 1.15)
    set_style(doc.styles["List Bullet"], 10.5, DARK, False, 0, 4, 1.25)
    set_style(doc.styles["List Number"], 10.5, DARK, False, 0, 4, 1.25)
    for style_name in ("List Bullet", "List Number"):
        pf = doc.styles[style_name].paragraph_format
        pf.left_indent = Inches(0.375)
        pf.first_line_indent = Inches(-0.188)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Chayapon Works  |  คู่มือการใช้งานเว็บไซต์")
    font_run(r, 8.5, MUTED, True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("คู่มือฉบับ 1.0  •  หน้า ")
    font_run(r, 8.5, MUTED)
    add_field(p, "PAGE")


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    font_run(r, 9, BLUE, True)
    r.font.letter_spacing = Pt(1) if hasattr(r.font, "letter_spacing") else None
    return p


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph(style="Title")
    p.add_run(text)
    if subtitle:
        p = doc.add_paragraph(style="Subtitle")
        p.add_run(subtitle)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        font_run(r, bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc, steps):
    for step in steps:
        doc.add_paragraph(step, style="List Number")


def add_callout(doc, title, text, tone="blue"):
    colors = {
        "blue": (LIGHT_BLUE, BLUE),
        "green": ("ECFDF3", GREEN),
        "amber": ("FFF4E5", AMBER),
        "red": ("FEF3F2", RED),
        "gray": (PALE, MUTED),
    }
    fill, accent = colors[tone]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), accent)
        borders.append(border)
    p_pr.append(borders)
    r = p.add_run(title)
    font_run(r, 10.5, accent, True)
    r = p.add_run("\n" + text)
    font_run(r, 9.5, DARK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        font_run(r, 9.5, BLUE, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            font_run(r, 9.2, DARK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)

    # Cover: editorial_cover pattern, branded for Chayapon Works.
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Chayapon")
    font_run(r, 15, BLUE, True)
    r = p.add_run(" Works")
    font_run(r, 15, DARK, True)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("คู่มือการใช้งานเว็บไซต์")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("สำหรับลูกค้า พนักงาน และผู้ดูแลระบบ")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(70)
    r = p.add_run("รถยนต์มือสอง  •  อสังหาริมทรัพย์  •  อะไหล่รถยนต์")
    font_run(r, 10.5, BLUE, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ฉบับ 1.0  |  สิงหาคม 2026")
    font_run(r, 10, MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ใช้กับเว็บไซต์แบบ Responsive ทั้งคอมพิวเตอร์ แท็บเล็ต และมือถือ")
    font_run(r, 9.5, MUTED, italic=True)

    page_break(doc)
    add_kicker(doc, "HOW TO USE THIS GUIDE")
    add_title(doc, "เริ่มต้นใช้งานคู่มือนี้", "เลือกอ่านเฉพาะส่วนที่ตรงกับบทบาทของคุณ")
    add_callout(doc, "ขอบเขตของคู่มือ", "คู่มือนี้อ้างอิงฟังก์ชันที่มีอยู่ในระบบจริง ณ เวอร์ชันปัจจุบัน ชื่อปุ่มหรือตำแหน่งอาจเปลี่ยนเล็กน้อยเมื่อมีการปรับปรุงเว็บไซต์", "blue")
    add_table(doc, ["บทบาท", "เหมาะสำหรับ", "บทที่ควรอ่าน"], [
        ["ลูกค้า (Customer)", "ผู้ค้นหารถ อสังหาริมทรัพย์ หรือซื้ออะไหล่", "บทที่ 2-7"],
        ["พนักงาน (Staff)", "ผู้ติดต่อลูกค้า ดูนัดหมาย และจัดการออเดอร์", "บทที่ 8-10"],
        ["ผู้ดูแลระบบ (Admin)", "เจ้าของร้านหรือผู้ดูแลสินค้า ผู้ใช้ และการตั้งค่าร้าน", "บทที่ 8-14"],
    ], [1800, 4300, 3260])
    doc.add_heading("สัญลักษณ์ที่ใช้ในคู่มือ", level=2)
    add_bullets(doc, [
        "ขั้นตอน: ลำดับการทำงานที่ควรทำตามทีละข้อ",
        "ข้อควรทราบ: เงื่อนไขหรือข้อมูลสำคัญก่อนดำเนินการ",
        "คำเตือน: การกระทำที่อาจกระทบข้อมูล สต็อก สถานะออเดอร์ หรือสิทธิ์ผู้ใช้",
        "ตรวจสอบผล: จุดที่ใช้ยืนยันว่างานสำเร็จแล้ว",
    ])
    doc.add_heading("สารบัญ", level=2)
    add_table(doc, ["สำหรับลูกค้า", "หลังบ้านและการดูแลระบบ"], [
        ["1. ภาพรวมระบบและสิทธิ์ใช้งาน", "8. การเข้าสู่ระบบหลังบ้าน"],
        ["2. หน้าแรกและการค้นหา", "9. งานประจำวันของ Staff"],
        ["3. รถยนต์และการนัดหมาย", "10. การจัดการออเดอร์และสลิป"],
        ["4. อสังหาริมทรัพย์ บ้าน คอนโด และที่ดิน", "11. การจัดการสินค้า"],
        ["5. อะไหล่ ตะกร้า และ Checkout", "12. การจัดการผู้ใช้และสิทธิ์"],
        ["6. บัญชีลูกค้าและการติดตามรายการ", "13. การตั้งค่าร้านและช่องทางชำระเงิน"],
        ["7. สมัครสมาชิกและรหัสผ่าน", "14. ความปลอดภัย การแก้ปัญหา และ Checklist"],
    ], [4680, 4680])

    add_kicker(doc, "SYSTEM OVERVIEW")
    doc.add_heading("1. ภาพรวมระบบและสิทธิ์ใช้งาน", level=1)
    add_body(doc, "Chayapon Works เป็นเว็บไซต์ E-Commerce ของผู้ขายรายเดียว แบ่งสินค้าเป็นรถยนต์มือสอง อสังหาริมทรัพย์ และอะไหล่รถยนต์ ลูกค้าสามารถค้นหา ดูรายละเอียด นัดหมาย และซื้ออะไหล่ได้จากระบบเดียว")
    add_table(doc, ["ความสามารถ", "ลูกค้า", "Staff", "Admin"], [
        ["ค้นหาและดูสินค้า", "ได้", "ได้", "ได้"],
        ["ส่งคำขอนัดหมาย", "ได้เมื่อเข้าสู่ระบบ", "ได้ในฐานะลูกค้า", "ได้ในฐานะลูกค้า"],
        ["ซื้ออะไหล่", "Guest หรือสมาชิก", "ได้", "ได้"],
        ["จัดการ Lead/นัดหมาย", "ดูของตนเอง", "ได้", "ได้"],
        ["ตรวจสลิปและอัปเดตออเดอร์", "ดูสถานะ", "ได้", "ได้"],
        ["เพิ่ม/แก้ไขสินค้า", "ไม่ได้", "ไม่ได้", "ได้"],
        ["จัดการสิทธิ์และตั้งค่าร้าน", "ไม่ได้", "ไม่ได้", "ได้"],
    ], [3300, 2020, 2020, 2020])
    add_callout(doc, "หลักการแบ่งสิทธิ์", "เมนูของ Staff และ Admin เป็นคนละชุด ไม่ใช่เมนูเดียวกันที่เพียงถูกปิดไว้ ผู้ใช้ทั่วไปไม่สามารถเข้าสู่ Dashboard หลังบ้านได้", "green")
    doc.add_heading("เส้นทางหลัก", level=2)
    add_table(doc, ["งาน", "เส้นทาง"], [
        ["หน้าแรก", "/"],
        ["หน้าสินค้า", "/vehicles, /real-estate, /parts"],
        ["ซื้ออะไหล่", "/cart, /checkout"],
        ["บัญชีของฉัน", "/profile"],
        ["หลังบ้าน", "/dashboard"],
    ], [2500, 6860])

    page_break(doc)
    add_kicker(doc, "CUSTOMER GUIDE")
    doc.add_heading("2. หน้าแรกและการค้นหา", level=1)
    doc.add_heading("ค้นหาจาก Hero Search", level=2)
    add_steps(doc, [
        "เปิดหน้าแรก แล้วเลือกแท็บ รถยนต์, อสังหาริมทรัพย์ หรืออะไหล่",
        "กรอกคำค้นหรือเลือกตัวกรองหลัก เช่น ยี่ห้อ รุ่น ประเภท ทำเล และช่วงราคา",
        "กด “ค้นหา” เพื่อเปิดหน้ารายการสินค้าพร้อมเงื่อนไขที่เลือก",
        "ใช้ตัวกรองด้านข้างบนคอมพิวเตอร์ หรือปุ่ม “ตัวกรอง” บนมือถือเพื่อปรับผลลัพธ์",
        "กดการ์ดสินค้า หรือปุ่ม “ดูรายละเอียด/ติดต่อผู้ขาย” เพื่อเปิดหน้าสินค้า",
    ])
    doc.add_heading("หมวดหมู่ที่รองรับ", level=2)
    add_table(doc, ["หมวด", "ตัวกรองสำคัญ", "การดำเนินการหลัก"], [
        ["รถยนต์มือสอง", "ยี่ห้อ รุ่น ปี ราคา ไมล์ เชื้อเพลิง ตัวถัง", "ดูรายละเอียดและนัดทดลองขับ"],
        ["อสังหาริมทรัพย์", "ประเภท ทำเล ราคา พื้นที่ ห้องนอน/ห้องน้ำ", "ดูรายละเอียดและนัดดูสถานที่"],
        ["อะไหล่รถยนต์", "คำค้น แบรนด์ หมวดหมู่ รุ่นรถ ราคา", "เพิ่มลงตะกร้าและซื้อทันที"],
    ], [1900, 4000, 3460])
    add_callout(doc, "บนมือถือ", "แถบเมนูและตัวกรองจะถูกย่อให้เหมาะกับหน้าจอ ใช้เมนูสามขีดหรือตัวกรองแบบแผงเลื่อน และตรวจสอบชิปเงื่อนไขด้านบนก่อนเลือกสินค้า", "gray")

    page_break(doc)
    doc.add_heading("3. รถยนต์และการนัดหมาย", level=1)
    doc.add_heading("ดูรายละเอียดรถ", level=2)
    add_bullets(doc, [
        "เลื่อนดูภาพทั้งหมดและตรวจสอบป้ายสถานะสินค้า",
        "ตรวจสอบราคา ปี ระยะทาง เชื้อเพลิง เกียร์ สี เครื่องยนต์ และสถานที่",
        "อ่านรายละเอียด จุดเด่น และข้อมูลการตรวจสภาพก่อนติดต่อร้าน",
        "ใช้ปุ่มโทรหรือ LINE เมื่อร้านตั้งค่าช่องทางติดต่อแล้ว",
    ])
    doc.add_heading("ส่งคำขอนัดดูรถ", level=2)
    add_steps(doc, [
        "เข้าสู่ระบบก่อนส่งคำขอ หากยังไม่ได้เข้าสู่ระบบ เว็บไซต์จะพาไปหน้า Sign in และกลับมาหน้ารถเดิมหลังสำเร็จ",
        "กด “ติดต่อร้าน / นัดหมายทดลองขับ”",
        "กรอกชื่อ เบอร์โทรศัพท์ และวันที่สะดวก (วันที่ไม่บังคับ)",
        "กด “ส่งข้อมูลติดต่อ” และรอข้อความยืนยัน",
        "ติดตามรายการได้ที่ บัญชีของฉัน > การนัดหมายของฉัน",
    ])
    add_callout(doc, "ข้อควรทราบ", "การส่งแบบฟอร์มคือคำขอให้ร้านติดต่อกลับ ยังไม่ถือเป็นการยืนยันเวลานัดจนกว่าพนักงานจะติดต่อยืนยัน", "amber")
    doc.add_heading("สถานะการนัดหมาย", level=2)
    add_table(doc, ["สถานะ", "ความหมาย", "สิ่งที่ลูกค้าควรทำ"], [
        ["ใหม่", "ร้านได้รับคำขอแล้ว", "รอการติดต่อกลับ"],
        ["ติดต่อแล้ว", "พนักงานติดต่อหรือกำลังประสานงาน", "ตรวจสอบโทรศัพท์/LINE"],
        ["ปิดแล้ว", "รายการได้รับการจัดการเสร็จสิ้น", "ติดต่อร้านหากต้องการนัดใหม่"],
    ], [1700, 4100, 3560])

    page_break(doc)
    doc.add_heading("4. อสังหาริมทรัพย์ บ้าน คอนโด และที่ดิน", level=1)
    add_body(doc, "หมวดอสังหาริมทรัพย์รองรับมากกว่าบ้าน โดยรวมบ้าน คอนโด ทาวน์เฮ้าส์ และที่ดิน ข้อมูลที่แสดงจะแตกต่างตามประเภททรัพย์")
    doc.add_heading("วิธีค้นหาและนัดดูสถานที่", level=2)
    add_steps(doc, [
        "เลือกแท็บ “อสังหาริมทรัพย์” จากหน้าแรก หรือเปิดเมนูอสังหาริมทรัพย์",
        "เลือกประเภททรัพย์ ทำเล ช่วงราคา และขนาดพื้นที่ตามต้องการ",
        "เปิดหน้ารายละเอียดเพื่อตรวจภาพ ราคา ที่ตั้ง พื้นที่ ห้อง และสิ่งอำนวยความสะดวก",
        "สำหรับที่ดิน ให้ตรวจขนาดพื้นที่ ทำเล ทางเข้า และรายละเอียดข้อจำกัดในประกาศ",
        "กด “ติดต่อร้าน / นัดดูสถานที่จริง” แล้วกรอกข้อมูลติดต่อ",
    ])
    add_callout(doc, "ตรวจสอบก่อนตัดสินใจ", "ข้อมูลบนเว็บไซต์ใช้สำหรับคัดกรองเบื้องต้น ควรนัดตรวจสถานที่จริง เอกสารสิทธิ์ ขอบเขตที่ดิน ภาระผูกพัน และค่าใช้จ่ายกับร้านก่อนทำธุรกรรม", "amber")
    doc.add_heading("ข้อมูลที่ควรเช็กตามประเภท", level=2)
    add_table(doc, ["ประเภท", "รายการตรวจสอบ"], [
        ["บ้าน/ทาวน์เฮ้าส์", "พื้นที่ใช้สอย ห้องนอน ห้องน้ำ ที่จอดรถ สภาพอาคาร และทำเล"],
        ["คอนโด", "ชั้น พื้นที่ ส่วนกลาง ที่จอดรถ เฟอร์นิเจอร์ และค่าใช้จ่ายส่วนกลาง"],
        ["ที่ดิน", "ขนาดพื้นที่ หน้ากว้าง ทางเข้า สาธารณูปโภค ผังเมือง และเอกสารสิทธิ์"],
    ], [2100, 7260])

    page_break(doc)
    doc.add_heading("5. อะไหล่ ตะกร้า และ Checkout", level=1)
    doc.add_heading("เพิ่มสินค้าในตะกร้า", level=2)
    add_steps(doc, [
        "เปิดหน้าอะไหล่และเลือกสินค้าที่ต้องการ",
        "ตรวจ SKU แบรนด์ หมวดหมู่ รุ่นรถที่รองรับ ราคา และสถานะสต็อก",
        "เลือกจำนวน แล้วกด “หยิบใส่ตะกร้า” หรือ “สั่งซื้อทันที” จากนั้นตรวจจำนวน ราคา และยอดรวม",
        "ปรับหรือลบสินค้าในตะกร้า แล้วกด “ดำเนินการชำระเงิน”",
    ])
    doc.add_heading("Checkout: 3 ขั้นตอน", level=2)
    add_table(doc, ["ขั้น", "ข้อมูลที่ต้องทำ", "ตรวจสอบก่อนต่อไป"], [
        ["1 ข้อมูลจัดส่ง", "ชื่อ เบอร์โทร อีเมล และที่อยู่", "จังหวัด เขต/อำเภอ แขวง/ตำบล รหัสไปรษณีย์ และรายละเอียดบ้านเลขที่"],
        ["2 ชำระเงิน", "สแกน PromptPay QR ตามยอดสุทธิ", "ชื่อบัญชี ยอดเงิน และเลขอ้างอิงคำสั่งซื้อ"],
        ["3 ยืนยัน/เสร็จสิ้น", "อัปโหลดสลิปและยืนยัน", "เลขคำสั่งซื้อและสถานะรอตรวจสลิป"],
    ], [1500, 3350, 4510])
    doc.add_heading("กรอกที่อยู่จัดส่ง", level=2)
    add_bullets(doc, [
        "เลือก จังหวัด > เขต/อำเภอ > แขวง/ตำบล > รหัสไปรษณีย์ ตามลำดับ",
        "กรอกบ้านเลขที่ ซอย หมู่ ถนน และรายละเอียดเพิ่มเติมด้วยตนเอง",
        "เบอร์โทรศัพท์จะแสดงรูปแบบอัตโนมัติ เช่น 081-234-5678",
        "สมาชิกที่เข้าสู่ระบบจะได้รับการเติมข้อมูลที่บันทึกไว้ ส่วน Guest Checkout ซื้อได้โดยไม่บังคับสมัครสมาชิก",
    ])
    doc.add_heading("ชำระผ่าน PromptPay และอัปโหลดสลิป", level=2)
    add_steps(doc, [
        "ตรวจ Thumbnail ชื่อสินค้า จำนวน ราคา ค่าจัดส่ง และยอดสุทธิในสรุปคำสั่งซื้อ",
        "สแกน QR Code และชำระยอดที่แสดงให้ตรงกัน",
        "กดพื้นที่อัปโหลดสลิป แล้วเลือกไฟล์ JPG, PNG หรือ WebP ที่เห็นข้อมูลชัดเจน",
        "กด “ยืนยันการชำระเงิน” เพียงครั้งเดียว รอระบบตอบกลับ และจดเลขคำสั่งซื้อไว้ติดตาม",
    ])
    add_callout(doc, "ระบบความปลอดภัยการชำระเงิน", "ออเดอร์จะอยู่สถานะ “รอตรวจสลิป” และยังไม่ถูกจัดส่งจนกว่า Staff หรือ Admin จะตรวจสอบและอนุมัติสลิป", "green")

    page_break(doc)
    doc.add_heading("6. บัญชีลูกค้าและการติดตามรายการ", level=1)
    doc.add_heading("My Account", level=2)
    add_bullets(doc, [
        "ภาพรวม: ดูคำสั่งซื้อล่าสุดและการนัดหมายล่าสุด",
        "คำสั่งซื้อของฉัน: ดูรายการ ราคา สถานะการชำระเงิน และสถานะจัดส่ง",
        "การนัดหมายของฉัน: ดูรถหรือทรัพย์ที่สนใจ สถานะ และปุ่มโทร/LINE",
        "ข้อมูลส่วนตัว: แก้ชื่อ เบอร์โทร LINE ที่อยู่ และรูปโปรไฟล์",
    ])
    doc.add_heading("สถานะคำสั่งซื้อ", level=2)
    add_table(doc, ["สถานะ", "ความหมาย"], [
        ["รอดำเนินการ", "ร้านได้รับคำสั่งซื้อและกำลังตรวจสอบ"],
        ["ยืนยันแล้ว", "ตรวจสอบการชำระเงินและรับออเดอร์แล้ว"],
        ["จัดส่งแล้ว", "สินค้าออกจากร้านแล้ว"],
        ["สำเร็จ", "การจัดส่งหรือรายการเสร็จสมบูรณ์"],
        ["ยกเลิก", "คำสั่งซื้อถูกยกเลิกและระบบคืนสต็อกตามเงื่อนไข"],
    ], [2400, 6960])
    add_callout(doc, "Guest Checkout", "ผู้ซื้อแบบ Guest ควรเก็บเลขคำสั่งซื้อและอีเมลที่ใช้สั่งซื้อไว้ ระบบจะเสนอให้สร้างบัญชีหลังทำรายการสำเร็จเพื่อช่วยติดตามครั้งถัดไป", "blue")

    doc.add_heading("7. สมัครสมาชิก เข้าสู่ระบบ และรหัสผ่าน", level=1)
    add_steps(doc, [
        "กด “เข้าสู่ระบบ” ที่แถบเมนู แล้วเลือกสมัครสมาชิกหากยังไม่มีบัญชี",
        "กรอกชื่อ อีเมล และรหัสผ่านอย่างน้อย 12 ตัวอักษร",
        "เข้าสู่ระบบด้วยอีเมลและรหัสผ่านที่สมัครไว้",
        "หากลืมรหัสผ่าน ให้กด “ลืมรหัสผ่าน” กรอกอีเมล และเปิดลิงก์รีเซ็ตจากอีเมล",
        "หลังเปลี่ยนรหัสผ่าน เซสชันเดิมจะถูกยกเลิกเพื่อความปลอดภัย",
    ])
    add_callout(doc, "คำแนะนำรหัสผ่าน", "ใช้รหัสผ่านที่ไม่ซ้ำกับบริการอื่น หลีกเลี่ยงชื่อ เบอร์โทร หรือวันเกิด และไม่ส่งรหัสผ่านผ่านแชตหรือ LINE", "red")

    page_break(doc)
    add_kicker(doc, "BACK OFFICE GUIDE")
    doc.add_heading("8. การเข้าสู่ระบบหลังบ้าน", level=1)
    add_body(doc, "หลังเข้าสู่ระบบ ผู้ใช้ที่มีบทบาท ADMIN หรือ STAFF สามารถเปิด /dashboard ได้ ระบบจะแสดง Sidebar ตามสิทธิ์จริง")
    add_table(doc, ["เมนู", "Staff", "Admin"], [
        ["งานวันนี้/ภาพรวม", "Today's Operations", "ภาพรวมธุรกิจ รายได้ และกิจกรรม"],
        ["ลูกค้าและการนัดหมาย", "เห็นและอัปเดต", "เห็นและอัปเดต"],
        ["คำสั่งซื้อ", "เห็นและอัปเดต", "เห็นและอัปเดต"],
        ["รายการสินค้า", "ไม่แสดง", "เพิ่ม แก้ไข และจัดการสถานะ"],
        ["ผู้ใช้งานและพนักงาน", "ไม่แสดง", "จัดการ Role"],
        ["ตั้งค่าร้านค้า", "ไม่แสดง", "แก้ไขข้อมูลติดต่อและการเงิน"],
    ], [3100, 3130, 3130])
    add_callout(doc, "ความปลอดภัย", "ห้ามให้ผู้อื่นใช้บัญชี Staff/Admin ร่วมกัน เพราะระบบบันทึก Audit Log ตามบัญชีผู้ดำเนินการ", "red")

    page_break(doc)
    doc.add_heading("9. งานประจำวันของ Staff", level=1)
    doc.add_heading("ลำดับงานแนะนำ", level=2)
    add_steps(doc, [
        "เปิด Dashboard และตรวจจำนวนงานรอติดต่อ นัดหมาย และออเดอร์รอตรวจ",
        "เปิด “ลูกค้าและการนัดหมาย” แล้วจัดการ Lead ใหม่ก่อน",
        "โทรหาลูกค้าจากเบอร์ในรายการ ยืนยันสินค้า วันที่ เวลา และรายละเอียดนัด",
        "อัปเดตสถานะจาก ใหม่ > ติดต่อแล้ว > ปิดแล้ว ให้ตรงกับงานจริง",
        "เปิด “คำสั่งซื้อ” เพื่อตรวจสลิปและอัปเดตสถานะการจัดส่ง",
    ])
    doc.add_heading("แนวทางการอัปเดต Lead", level=2)
    add_table(doc, ["สถานะ", "ใช้เมื่อ", "การดำเนินการถัดไป"], [
        ["ใหม่", "ยังไม่ติดต่อ", "โทรกลับตามลำดับความเร่งด่วน"],
        ["ติดต่อแล้ว", "คุยกับลูกค้าแล้วหรือกำลังรอยืนยัน", "บันทึกผลและติดตามตามเวลาที่ตกลง"],
        ["ปิดแล้ว", "นัดหมายเสร็จ ยกเลิก หรือไม่ต้องติดตามต่อ", "ตรวจว่าข้อมูลสุดท้ายครบ"],
    ], [1700, 4100, 3560])

    page_break(doc)
    doc.add_heading("10. การจัดการออเดอร์และสลิป", level=1)
    doc.add_heading("ตรวจและอนุมัติการชำระเงิน", level=2)
    add_steps(doc, [
        "เปิด Dashboard > คำสั่งซื้อ แล้วเลือกรายการที่สถานะ “รอตรวจสลิป”",
        "กด “ดูสลิป” ตรวจยอดเงิน วันเวลา เลขอ้างอิง และบัญชีปลายทาง",
        "ตรวจรายการสินค้า จำนวน ที่อยู่จัดส่ง และยอดสุทธิประกอบกัน",
        "หากข้อมูลถูกต้อง กด “อนุมัติ” ระบบจะตั้งการชำระเงินเป็น “ชำระแล้ว” และยืนยันออเดอร์ที่ยังรอดำเนินการ",
        "หากไม่ถูกต้อง กด “ไม่ผ่าน” แล้วติดต่อลูกค้าเพื่อขอสลิปใหม่หรือแก้ไขข้อมูล",
    ])
    add_callout(doc, "คำเตือน", "ห้ามเปลี่ยนสถานะเป็น ยืนยันแล้ว, จัดส่งแล้ว หรือสำเร็จ ก่อนการชำระเงินเป็น “ชำระแล้ว” ระบบจะปฏิเสธการดำเนินการดังกล่าว", "red")
    doc.add_heading("เปลี่ยนสถานะออเดอร์", level=2)
    add_table(doc, ["จาก", "ไปยัง", "เงื่อนไข/ผลลัพธ์"], [
        ["รอดำเนินการ", "ยืนยันแล้ว", "ต้องตรวจสลิปผ่าน"],
        ["ยืนยันแล้ว", "จัดส่งแล้ว", "แพ็กสินค้าและมีข้อมูลจัดส่งพร้อม"],
        ["จัดส่งแล้ว", "สำเร็จ", "ส่งมอบสำเร็จ"],
        ["สถานะใด ๆ ที่อนุญาต", "ยกเลิก", "ระบบคืนสต็อกหนึ่งครั้งและบันทึก Audit Log"],
    ], [2200, 2200, 4960])

    page_break(doc)
    doc.add_heading("11. การจัดการสินค้า", level=1)
    doc.add_heading("เพิ่มสินค้าใหม่", level=2)
    add_steps(doc, [
        "เข้าสู่ระบบด้วย Admin แล้วเปิด รายการสินค้า > เพิ่มสินค้าใหม่",
        "เลือกประเภท รถยนต์, อสังหาริมทรัพย์ หรืออะไหล่ ประเภทจะเปลี่ยนชุดช่องข้อมูลอัตโนมัติ",
        "กรอกชื่อ ราคา รายละเอียด สถานะ และเลือก “สินค้าแนะนำ” เมื่อต้องการแสดงเด่นหน้าแรก",
        "เพิ่มรูปภาพสูงสุด 8 รูป ไฟล์ต้นฉบับไม่เกิน 5 MB ต่อรูป และให้ขนาดรวมหลังแปลงไม่เกิน 7 MB",
        "กรอกข้อมูลเฉพาะประเภทให้ครบ แล้วกดบันทึก",
        "เปิดหน้าร้านตรวจชื่อ ราคา รูป และรายละเอียดหลังบันทึกทุกครั้ง",
    ])
    doc.add_heading("ข้อมูลเฉพาะประเภท", level=2)
    add_table(doc, ["ประเภท", "ข้อมูลสำคัญ"], [
        ["รถยนต์", "ยี่ห้อ รุ่น ปี ไมล์ เกียร์ เชื้อเพลิง สี ตัวถัง ระบบขับเคลื่อน จังหวัดป้ายทะเบียน"],
        ["อสังหาริมทรัพย์", "ประเภททรัพย์ พื้นที่ ห้องนอน ห้องน้ำ ทำเล เขต สถานที่ใกล้เคียง และเฟอร์นิเจอร์"],
        ["อะไหล่", "SKU แบรนด์ หมวดหมู่ รุ่นรถที่รองรับ และจำนวนสต็อก"],
    ], [1900, 7460])
    doc.add_heading("สถานะสินค้า", level=2)
    add_bullets(doc, [
        "Active: เปิดขายและแสดงบนหน้าร้าน",
        "Reserved: จองแล้ว ควรตรวจว่าการ์ดและรายละเอียดสื่อสารตรงกัน",
        "Sold: ขายแล้ว ไม่ควรรับคำขอใหม่",
        "การลบสินค้าเป็นการกระทำสำคัญ ควรตรวจว่าสินค้าไม่มีรายการที่ยังดำเนินการอยู่",
    ])

    page_break(doc)
    doc.add_heading("12. การจัดการผู้ใช้และสิทธิ์", level=1)
    add_body(doc, "เมนูผู้ใช้งานและพนักงานมีเฉพาะ Admin ใช้ดู Role สถานะกิจกรรมล่าสุด และเปลี่ยนสิทธิ์บัญชี")
    doc.add_heading("เปลี่ยน Role", level=2)
    add_steps(doc, [
        "ค้นหาผู้ใช้จากชื่อหรืออีเมลในหน้า ผู้ใช้งานและพนักงาน",
        "เปิดเมนูจัดการของบัญชีที่ต้องการ แล้วเลือก Role ใหม่",
        "อ่านรายละเอียดผลกระทบในหน้าต่างยืนยัน",
        "กดยืนยัน ระบบจะเปลี่ยนสิทธิ์และบันทึกผู้ดำเนินการ เวลา ค่าเดิม และค่าใหม่ใน Audit Log",
    ])
    add_table(doc, ["Role", "สิทธิ์หลัก"], [
        ["USER", "ใช้หน้าร้าน บัญชีลูกค้า คำสั่งซื้อ และนัดหมายของตนเอง"],
        ["STAFF", "ใช้ Dashboard งานวันนี้ จัดการ Lead และออเดอร์"],
        ["ADMIN", "เข้าถึงงานทั้งหมด รวมสินค้า ผู้ใช้ การตั้งค่าร้าน และข้อมูลภาพรวมธุรกิจ"],
    ], [1800, 7560])
    add_callout(doc, "ข้อจำกัดสำคัญ", "ระบบไม่อนุญาตให้ Admin ลดสิทธิ์บัญชีของตนเอง เพื่อลดความเสี่ยงที่ระบบจะไม่มีผู้ดูแล ควรตรวจอีเมลและชื่อบัญชีก่อนเปลี่ยน Role ทุกครั้ง", "red")

    page_break(doc)
    doc.add_heading("13. การตั้งค่าร้านและช่องทางชำระเงิน", level=1)
    add_body(doc, "ข้อมูลติดต่อและการรับชำระเงินตั้งค่าจาก Dashboard โดย Admin ไม่ควรแก้ไฟล์ Environment สำหรับข้อมูลที่เปลี่ยนบ่อย")
    doc.add_heading("ข้อมูลที่ตั้งค่าได้", level=2)
    add_table(doc, ["กลุ่ม", "รายการ", "จุดที่นำไปใช้"], [
        ["ช่องทางติดต่อ", "ชื่อร้าน เบอร์โทร LINE URL", "หน้าสินค้า Modal ติดต่อ และหน้าการนัดหมาย"],
        ["PromptPay", "หมายเลข PromptPay และชื่อบัญชี", "QR Checkout และข้อความยืนยันยอด"],
        ["บัญชีธนาคาร", "ธนาคาร ชื่อบัญชี เลขบัญชี", "ข้อมูลอ้างอิงการชำระเงินตามที่ระบบรองรับ"],
    ], [1900, 3600, 3860])
    doc.add_heading("วิธีบันทึกการตั้งค่า", level=2)
    add_steps(doc, [
        "เปิด Dashboard > ตั้งค่าร้านค้า",
        "กรอกเบอร์โทร 10 หลัก และ LINE URL แบบเต็ม เช่น https://line.me/R/ti/p/@ชื่อบัญชี",
        "กรอก PromptPay เป็นเบอร์โทร 10 หลักหรือเลขนิติบุคคล 13 หลัก พร้อมชื่อบัญชี",
        "กดแสดงข้อมูลการเงินเมื่อต้องตรวจค่า Sensitive และซ่อนอีกครั้งเมื่อเสร็จ",
        "กด “บันทึกการตั้งค่า” อ่านหน้าต่างยืนยัน แล้วกดยืนยัน",
        "ตรวจข้อความ Success และทดลองเปิดหน้าสินค้า/Checkout เพื่อยืนยันว่าปุ่มโทร LINE และ QR ทำงาน",
    ])
    add_callout(doc, "การปกป้องข้อมูล", "หมายเลข PromptPay และเลขบัญชีถูกเข้ารหัสก่อนบันทึก Public API ไม่ส่งค่าฉบับเต็ม และการแก้ไขทุกครั้งถูกบันทึกใน Audit Log", "green")
    add_callout(doc, "หาก Checkout แจ้งว่ายังไม่ได้ตั้ง PromptPay", "ให้ Admin กรอกหมายเลข PromptPay และชื่อบัญชีแล้วบันทึกใหม่ จากนั้นเปิด Checkout อีกครั้ง ไม่ควรแก้ฐานข้อมูลโดยตรง", "amber")

    page_break(doc)
    doc.add_heading("14. ความปลอดภัย การแก้ปัญหา และ Checklist", level=1)
    doc.add_heading("แนวปฏิบัติสำหรับทุกบทบาท", level=2)
    add_bullets(doc, [
        "ตรวจโดเมนและสัญลักษณ์ HTTPS ก่อนกรอกรหัสผ่านหรือข้อมูลการเงิน",
        "ไม่แชร์รหัสผ่าน โค้ดรีเซ็ต หรือสลิปที่มีข้อมูลสำคัญในช่องทางสาธารณะ",
        "ออกจากระบบเมื่อใช้เครื่องร่วม โดยเฉพาะบัญชี Staff/Admin",
        "หากพบการเปลี่ยนแปลงผิดปกติ ให้หยุดดำเนินการและแจ้ง Admin ตรวจ Audit Log",
        "หลีกเลี่ยงการกดปุ่มชำระเงินหรือบันทึกซ้ำขณะระบบกำลังประมวลผล",
    ])
    doc.add_heading("แก้ปัญหาที่พบบ่อย", level=2)
    add_table(doc, ["อาการ", "สาเหตุที่เป็นไปได้", "วิธีแก้"], [
        ["เปิด Dashboard ไม่ได้", "ยังไม่เข้าสู่ระบบหรือไม่มี Role", "เข้าสู่ระบบด้วยบัญชี Staff/Admin หรือติดต่อ Admin"],
        ["ไม่เห็นปุ่ม LINE/โทร", "Admin ยังไม่ตั้งค่าช่องทาง", "Admin บันทึกเบอร์โทรและ LINE URL ในตั้งค่าร้าน"],
        ["สร้าง QR ไม่ได้", "ยังไม่มี PromptPay หรือข้อมูลไม่ถูกต้อง", "Admin ตรวจหมายเลข 10/13 หลักและชื่อบัญชี"],
        ["อัปโหลดรูป/สลิปไม่ได้", "ชนิดไฟล์หรือขนาดไม่รองรับ", "ใช้ JPG, PNG หรือ WebP และลดขนาดไฟล์"],
        ["ยืนยันออเดอร์ไม่ได้", "สลิปยังไม่ผ่าน", "ตรวจสลิปแล้วกดอนุมัติก่อนเปลี่ยนสถานะ"],
        ["ไม่ได้อีเมลรีเซ็ตรหัสผ่าน", "อีเมลผิด สแปม หรือระบบส่งอีเมลยังไม่ตั้งค่า", "ตรวจ Spam แล้วให้ Admin ตรวจ Resend configuration"],
        ["ระบบแจ้งทำรายการบ่อยเกินไป", "Rate Limit ป้องกันการยิงซ้ำ", "รอตามเวลาที่แจ้งและหลีกเลี่ยงการกดซ้ำ"],
    ], [2400, 3260, 3700])
    page_break(doc)
    doc.add_heading("Checklist เปิดร้านประจำวัน", level=2)
    add_bullets(doc, [
        "ตรวจหน้าแรกและสินค้าแนะนำ",
        "ตรวจ Lead ใหม่และนัดหมายวันนี้",
        "ตรวจออเดอร์รอตรวจสลิป",
        "ตรวจสต็อกอะไหล่และสถานะสินค้าที่ขาย/จองแล้ว",
        "ทดสอบเบอร์โทร LINE และ PromptPay เมื่อมีการเปลี่ยนข้อมูล",
        "ตรวจ Activity/Audit Log เมื่อพบเหตุผิดปกติ",
    ])
    doc.add_heading("Checklist ก่อนปิดงาน", level=2)
    add_bullets(doc, [
        "Lead ที่ติดต่อแล้วมีสถานะถูกต้อง",
        "ออเดอร์ที่อนุมัติมีสลิปและยอดตรงกัน",
        "ออเดอร์จัดส่งแล้วมีสถานะล่าสุด",
        "ไม่มีบัญชี Staff/Admin เปิดค้างบนเครื่องสาธารณะ",
        "เหตุผิดปกติถูกบันทึกและแจ้งผู้รับผิดชอบแล้ว",
    ])
    add_callout(doc, "ช่องทางขอความช่วยเหลือ", "ใช้เบอร์โทรหรือ LINE ที่ Admin ตั้งค่าไว้ในเว็บไซต์ หากปัญหาเกี่ยวข้องกับสิทธิ์หรือข้อมูลหลังบ้าน ให้แจ้งชื่อบัญชี เวลาที่พบปัญหา หน้าที่ใช้งาน และข้อความผิดพลาด โดยไม่ส่งรหัสผ่าน", "blue")

    # Document properties and save.
    doc.core_properties.title = "คู่มือการใช้งานเว็บไซต์ Chayapon Works"
    doc.core_properties.subject = "คู่มือสำหรับ Customer, Staff และ Admin"
    doc.core_properties.author = "Chayapon Works"
    doc.core_properties.keywords = "Chayapon Works, คู่มือ, E-Commerce, รถยนต์, อสังหาริมทรัพย์, อะไหล่"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
